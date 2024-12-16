from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

class DOCXManager:

    def __init__(self, analysis_mode : str = "file"):
        self.doc = Document()
        self.sections = self.doc.sections
        self.doc_style = self.doc.styles["Normal"]
        font = self.doc_style.font
        font.name = "Arial"
        font.size = Pt(10)
        self.add_heading()
        

    def add_heading(self):
        self.doc.add_heading("Automated VirusTotal Report", 1)
        timestamp = self.doc.add_paragraph().add_run(f"Generated on {datetime.now().strftime('%d %B, %Y at %I:%M %p')}")
        timestamp.italic = True
        self.doc.add_paragraph("\n")

    def add_report(self, data):

        # Add fields
        fields = [
            ("SHA256 Hash", data.get("file_hash", "N/A")),
            ("Analysis", data.get("analysis", "N/A")),
            ("Reputation", str(data.get("community_score", "N/A"))),
            ("File Type", str(data.get("file_type", "N/A"))),
            ("File Size", str(data.get("file_size", "N/A"))),
            ("File Name", data.get("file_name", "N/A")),
            ("Last Analysis Date", data.get("file_last_analysis", "N/A")),
        ]

        for label, value in fields:
                paragraph = self.doc.add_paragraph()
                paragraph.add_run(f"{label}: ").bold = True
                paragraph.add_run(value)
                paragraph.style = self.doc.styles["Normal"]
                
        paragraph.add_run("\n")
        self.insertHR(paragraph)
            
    def configure_margin(self):
        for section in self.sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)

    def insertHR(self, paragraph):
        # https://stackoverflow.com/a/68530806
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
            'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
            'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
            'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange'
        )
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)

    def save(self, path):
        self.configure_margin()
        self.doc.save(path)

